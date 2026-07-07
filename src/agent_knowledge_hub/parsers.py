from __future__ import annotations

import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from agent_knowledge_hub.utils import normalize_space, read_text_with_fallback


class DocumentParseError(RuntimeError):
    """Raised when a document cannot be parsed into canonical blocks."""


class UnsupportedDocumentFormatError(DocumentParseError):
    """Raised when the current ingestion pipeline does not support a file type."""


@dataclass(frozen=True)
class ParsedBlock:
    block_type: str
    text: str
    page_start: int | None = None
    page_end: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    source_format: str
    parser_name: str
    page_count: int | None
    blocks: list[ParsedBlock]
    warnings: list[str] = field(default_factory=list)
    quality_report: dict[str, Any] | None = None


@dataclass(frozen=True)
class PdfTextQualityReport:
    char_count: int
    cjk_count: int
    mojibake_suspect_count: int
    mojibake_per_1k_cjk: float
    latin1_mojibake_ratio: float
    reason_codes: list[str]

    @property
    def should_use_ocr(self) -> bool:
        return bool(self.reason_codes)

    def to_warning(self) -> str:
        reasons = ",".join(self.reason_codes) if self.reason_codes else "none"
        return (
            "pdf_text_quality:"
            f" char_count={self.char_count};"
            f" cjk_count={self.cjk_count};"
            f" mojibake_suspect_count={self.mojibake_suspect_count};"
            f" mojibake_per_1k_cjk={self.mojibake_per_1k_cjk:.2f};"
            f" reason_codes={reasons}"
        )

    def to_dict(self, *, fallback_used: bool, fallback_parser: str | None) -> dict[str, Any]:
        score = 100.0
        if "text_too_short" in self.reason_codes:
            score -= 45.0
        if "mojibake_suspect_ratio_high" in self.reason_codes:
            score -= min(55.0, self.mojibake_per_1k_cjk / 4.0)
        if "latin1_mojibake_ratio_high" in self.reason_codes:
            score -= min(50.0, self.latin1_mojibake_ratio * 250.0)
        if fallback_used:
            score += 18.0
        score = max(0.0, min(100.0, score))
        status = "ok"
        if self.reason_codes and not fallback_used:
            status = "low_quality"
        elif self.reason_codes and fallback_used:
            status = "recovered_by_fallback"

        return {
            "score": round(score, 2),
            "status": status,
            "fallback_used": fallback_used,
            "fallback_parser": fallback_parser,
            "reason_codes": list(self.reason_codes),
            "metrics": {
                "char_count": self.char_count,
                "cjk_count": self.cjk_count,
                "mojibake_suspect_count": self.mojibake_suspect_count,
                "mojibake_per_1k_cjk": round(self.mojibake_per_1k_cjk, 2),
                "latin1_mojibake_ratio": round(self.latin1_mojibake_ratio, 4),
            },
        }


SUPPORTED_EXTENSIONS = {
    ".md",
    ".markdown",
    ".txt",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tiff",
    ".tif",
    ".gif",
    ".webp",
}

_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".gif", ".webp"}

_MIN_TRUSTED_PDF_TEXT_CHARS = 40
_MIN_CJK_CHARS_FOR_MOJIBAKE_CHECK = 20
_MIN_MOJIBAKE_SUSPECT_CHARS = 3
_MAX_MOJIBAKE_PER_1K_CJK = 50.0
_MAX_LATIN1_MOJIBAKE_RATIO = 0.12
_RAPIDOCR_RENDER_SCALE = 2.0
_RAPIDOCR_TEXT_SCORE = 0.45
_LATIN1_MOJIBAKE_CHARS = set(
    "ÃÂÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞß"
    "àáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ"
    "¡¢£¤¥¦§¨ª«¬®¯±²³´µ¶¸¹º»¼½¾¿"
)


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return parse_markdown(path)
    if suffix == ".txt":
        return parse_text(path)
    if suffix in {".html", ".htm"}:
        return parse_html(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix in _IMAGE_EXTENSIONS:
        return parse_image(path)
    raise UnsupportedDocumentFormatError(
        f"Unsupported document format '{suffix}' for {path}. "
        f"Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def parse_markdown(path: Path) -> ParsedDocument:
    text = read_text_with_fallback(path)
    blocks = _parse_markdown_blocks(text)
    return ParsedDocument(
        source_format="markdown",
        parser_name="stdlib-markdown-block-parser",
        page_count=None,
        blocks=blocks,
        quality_report=_build_generic_quality_report(
            blocks=blocks,
            source_format="markdown",
        ),
    )


def parse_text(path: Path) -> ParsedDocument:
    text = read_text_with_fallback(path)
    blocks = [
        ParsedBlock(block_type="paragraph", text=paragraph)
        for paragraph in _split_paragraphs(text)
    ]
    return ParsedDocument(
        source_format="text",
        parser_name="stdlib-text-block-parser",
        page_count=None,
        blocks=blocks,
        quality_report=_build_generic_quality_report(
            blocks=blocks,
            source_format="text",
        ),
    )


def parse_html(path: Path) -> ParsedDocument:
    text = read_text_with_fallback(path)
    parser = _BlockHtmlParser()
    parser.feed(text)
    return ParsedDocument(
        source_format="html",
        parser_name="stdlib-html-parser",
        page_count=None,
        blocks=parser.blocks,
        quality_report=_build_generic_quality_report(
            blocks=parser.blocks,
            source_format="html",
        ),
    )


def parse_pdf(path: Path) -> ParsedDocument:
    page_texts, page_count, warnings = _extract_pdf_text_pages_with_pypdf(path)
    blocks = _build_pdf_text_layer_blocks(page_texts)
    if not blocks:
        warnings.append("no_text_blocks_extracted")

    quality_report = assess_pdf_text_quality("\n".join(page_texts))
    if quality_report.should_use_ocr:
        warnings.append(quality_report.to_warning())
        try:
            ocr_document = _parse_pdf_with_rapidocr(path)
        except Exception as exc:  # pragma: no cover - exact failures depend on local optional deps
            warnings.append("pdf_text_quality_low_ocr_unavailable")
            warnings.append(f"pdf_ocr_fallback_failed: {exc}")
            return ParsedDocument(
                source_format="pdf",
                parser_name="pypdf",
                page_count=page_count,
                blocks=blocks,
                warnings=warnings,
                quality_report=quality_report.to_dict(
                    fallback_used=False,
                    fallback_parser=None,
                ),
            )

        if not _parsed_document_has_text_blocks(ocr_document):
            return ParsedDocument(
                source_format="pdf",
                parser_name="pypdf",
                page_count=page_count,
                blocks=blocks,
                warnings=[
                    *warnings,
                    "pdf_text_quality_low_ocr_unusable",
                    *ocr_document.warnings,
                ],
                quality_report=quality_report.to_dict(
                    fallback_used=False,
                    fallback_parser=None,
                ),
            )

        return ParsedDocument(
            source_format="pdf",
            parser_name="pypdf+rapidocr",
            page_count=ocr_document.page_count or page_count,
            blocks=ocr_document.blocks,
            warnings=[
                *warnings,
                "pdf_text_quality_low_using_ocr",
                *ocr_document.warnings,
            ],
            quality_report=quality_report.to_dict(
                fallback_used=True,
                fallback_parser="rapidocr",
            ),
        )

    return ParsedDocument(
        source_format="pdf",
        parser_name="pypdf",
        page_count=page_count,
        blocks=blocks,
        warnings=warnings,
        quality_report=quality_report.to_dict(
            fallback_used=False,
            fallback_parser=None,
        ),
    )


def assess_pdf_text_quality(text: str) -> PdfTextQualityReport:
    normalized_text = normalize_space(text)
    char_count = len(normalized_text)
    cjk_count = sum(1 for char in normalized_text if _is_cjk_char(char))
    mojibake_suspect_count = sum(
        1 for char in normalized_text if _is_mojibake_suspect_char(char)
    )
    mojibake_per_1k_cjk = (
        mojibake_suspect_count * 1000 / cjk_count if cjk_count else 0.0
    )
    latin1_mojibake_ratio = (
        sum(1 for char in normalized_text if char in _LATIN1_MOJIBAKE_CHARS) / char_count
        if char_count
        else 0.0
    )

    reason_codes: list[str] = []
    if char_count < _MIN_TRUSTED_PDF_TEXT_CHARS:
        reason_codes.append("text_too_short")
    if (
        cjk_count >= _MIN_CJK_CHARS_FOR_MOJIBAKE_CHECK
        and mojibake_suspect_count >= _MIN_MOJIBAKE_SUSPECT_CHARS
        and mojibake_per_1k_cjk > _MAX_MOJIBAKE_PER_1K_CJK
    ):
        reason_codes.append("mojibake_suspect_ratio_high")
    if (
        char_count >= _MIN_TRUSTED_PDF_TEXT_CHARS
        and cjk_count < _MIN_CJK_CHARS_FOR_MOJIBAKE_CHECK
        and latin1_mojibake_ratio > _MAX_LATIN1_MOJIBAKE_RATIO
    ):
        reason_codes.append("latin1_mojibake_ratio_high")

    return PdfTextQualityReport(
        char_count=char_count,
        cjk_count=cjk_count,
        mojibake_suspect_count=mojibake_suspect_count,
        mojibake_per_1k_cjk=mojibake_per_1k_cjk,
        latin1_mojibake_ratio=latin1_mojibake_ratio,
        reason_codes=reason_codes,
    )


def _extract_pdf_text_pages_with_pypdf(path: Path) -> tuple[list[str], int, list[str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError(
            "PDF parsing requires the optional dependency 'pypdf'. "
            "Install it before ingesting PDF files."
        ) from exc

    reader = PdfReader(str(path))
    page_texts: list[str] = []
    warnings: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - depends on PDF internals
            warnings.append(f"page {page_index}: text extraction failed: {exc}")
            text = ""
        page_texts.append(text)

    return page_texts, len(reader.pages), warnings


def _build_pdf_text_layer_blocks(page_texts: list[str]) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    seen_heading_keys: set[str] = set()
    for page_index, text in enumerate(page_texts, start=1):
        paragraphs = _split_pdf_text_paragraphs(text)
        for paragraph_index, paragraph in enumerate(paragraphs):
            heading_level = _classify_pdf_text_heading(paragraph)
            if heading_level is not None:
                heading_key = _pdf_heading_key(paragraph)
                if (
                    page_index > 1
                    and (paragraph_index == 0 or paragraph_index == len(paragraphs) - 1)
                    and heading_level > 1
                    and heading_key in seen_heading_keys
                ):
                    continue
                seen_heading_keys.add(heading_key)
                blocks.append(
                    ParsedBlock(
                        block_type="heading",
                        text=paragraph,
                        page_start=page_index,
                        page_end=page_index,
                        metadata={"level": heading_level, "pdf_text_heading": True},
                    )
                )
            else:
                blocks.append(
                    ParsedBlock(
                        block_type="paragraph",
                        text=paragraph,
                        page_start=page_index,
                        page_end=page_index,
                    )
                )
    return blocks


def _split_pdf_text_paragraphs(text: str) -> list[str]:
    paragraphs: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        paragraph = normalize_space(" ".join(paragraph_lines))
        if paragraph:
            paragraphs.append(paragraph)
        paragraph_lines = []

    for raw_line in text.splitlines():
        line = normalize_space(raw_line)
        if not line:
            flush_paragraph()
            continue
        if _is_pdf_text_noise_line(line):
            continue
        if _classify_pdf_text_heading(line) is not None:
            flush_paragraph()
            paragraphs.append(line)
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    return paragraphs


def _classify_pdf_text_heading(text: str) -> int | None:
    normalized = normalize_space(text)
    if not normalized:
        return None
    if len(normalized) > 96:
        return None
    if _is_non_heading_pdf_text_line(normalized):
        return None
    if normalized.endswith((".", ";", ":", ",")):
        return None

    if re.fullmatch(r"(?i)chapter\s+\d+[a-z]?", normalized):
        return 1
    if re.fullmatch(r"\d+(?:\.\d+){1,4}\s+\S.{2,}", normalized):
        if re.fullmatch(r"\d+\.\s+\S.{2,}", normalized):
            return None
        depth = normalized.split(" ", 1)[0].count(".") + 1
        return max(1, min(depth, 6))

    words = normalized.split()
    if 2 <= len(words) <= 8 and _looks_like_title_case_heading(words):
        return 2
    if 2 <= len(words) <= 8 and _looks_like_sentence_case_heading(words):
        return 2
    return None


def _is_pdf_text_noise_line(text: str) -> bool:
    normalized = normalize_space(text)
    if not normalized:
        return True
    if re.fullmatch(r"\d{1,4}", normalized):
        return True
    if re.fullmatch(r"(?i)copyright\s+.*(?:\b|[a-z])\d{1,4}$", normalized):
        return True
    return False


def _is_non_heading_pdf_text_line(text: str) -> bool:
    if text[0] in {"•", "*", "#"}:
        return True
    if re.match(r"^\d+[\.)]\s+\S", text):
        return True
    if re.match(r"^[A-Za-z]+:\s", text):
        return True
    if re.search(r"https?://|www\.|@", text, re.IGNORECASE):
        return True
    if any(char in text for char in {"=", "|", "\\"}):
        return True
    if text.count("/") >= 2:
        return True
    if re.search(r"\s-{1,2}[\w-]+", text):
        return True
    if re.search(r"\b[A-Za-z0-9_]*_[A-Za-z0-9_]*\b", text):
        return True
    return False


def _pdf_heading_key(text: str) -> str:
    return normalize_space(text).casefold()


def _looks_like_title_case_heading(words: list[str]) -> bool:
    strong_words = 0
    for word in words:
        stripped = word.strip("()[]{}")
        if not stripped:
            continue
        if stripped.lower() in {"and", "or", "of", "the", "a", "an", "to", "in", "at"}:
            continue
        if stripped[0].isupper() or stripped.isupper() or any(char.isdigit() for char in stripped):
            strong_words += 1
    return strong_words >= max(1, len(words) // 2)


def _looks_like_sentence_case_heading(words: list[str]) -> bool:
    first = words[0].strip("()[]{}")
    if not first or not first[0].isupper():
        return False
    lower_words = {word.strip("()[]{}").lower() for word in words}
    sentence_verbs = {
        "is",
        "are",
        "was",
        "were",
        "be",
        "been",
        "being",
        "can",
        "could",
        "must",
        "should",
        "shall",
        "will",
        "would",
        "do",
        "does",
        "did",
        "has",
        "have",
        "had",
        "describes",
        "requires",
        "contains",
        "provides",
        "uses",
    }
    return not bool(lower_words & sentence_verbs)


def _parse_pdf_with_rapidocr(path: Path) -> ParsedDocument:
    try:
        import fitz
        from rapidocr import RapidOCR
    except ImportError as exc:
        raise DocumentParseError(
            "PDF OCR fallback requires optional dependencies 'pymupdf', "
            "'rapidocr', and 'onnxruntime'. Install them before OCR fallback."
        ) from exc

    ocr = RapidOCR()
    document = fitz.open(str(path))
    page_count = document.page_count
    blocks: list[ParsedBlock] = []
    warnings: list[str] = []

    try:
        for page_index in range(page_count):
            page_number = page_index + 1
            try:
                page = document.load_page(page_index)
                pixmap = page.get_pixmap(
                    matrix=fitz.Matrix(_RAPIDOCR_RENDER_SCALE, _RAPIDOCR_RENDER_SCALE),
                    alpha=False,
                )
                result = ocr(pixmap.tobytes("png"), text_score=_RAPIDOCR_TEXT_SCORE)
                lines = _rapidocr_result_to_lines(result)
            except Exception as exc:  # pragma: no cover - depends on PDF/OCR internals
                warnings.append(f"page {page_number}: ocr failed: {exc}")
                lines = []

            blocks.append(
                ParsedBlock(
                    block_type="heading",
                    text=f"Page {page_number}",
                    page_start=page_number,
                    page_end=page_number,
                    metadata={"level": 2, "ocr_page_marker": True},
                )
            )
            if lines:
                blocks.append(
                    ParsedBlock(
                        block_type="paragraph",
                        text="\n".join(lines),
                        page_start=page_number,
                        page_end=page_number,
                        metadata={"ocr": True},
                    )
                )
            else:
                warnings.append(f"page {page_number}: no_ocr_text")
    finally:
        document.close()

    if not any(block.block_type != "heading" for block in blocks):
        warnings.append("no_ocr_text_blocks_extracted")

    return ParsedDocument(
        source_format="pdf",
        parser_name="rapidocr",
        page_count=page_count,
        blocks=blocks,
        warnings=warnings,
        quality_report=_build_generic_quality_report(
            blocks=blocks,
            source_format="pdf",
            fallback_parser="rapidocr",
        ),
    )


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise DocumentParseError(
            "DOCX parsing requires the optional dependency 'python-docx'. "
            "Install it before ingesting DOCX files."
        ) from exc

    document = DocxDocument(str(path))
    blocks: list[ParsedBlock] = []

    for paragraph in document.paragraphs:
        text = normalize_space(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        heading_match = re.search(r"heading\s+(\d+)", style_name)
        if heading_match:
            blocks.append(
                ParsedBlock(
                    block_type="heading",
                    text=text,
                    metadata={"level": int(heading_match.group(1))},
                )
            )
        else:
            blocks.append(ParsedBlock(block_type="paragraph", text=text))

    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [normalize_space(cell.text) for cell in row.cells]
            rows.append(" | ".join(cell for cell in cells if cell))
        table_text = "\n".join(row for row in rows if row)
        if table_text:
            blocks.append(ParsedBlock(block_type="table", text=table_text))

    return ParsedDocument(
        source_format="docx",
        parser_name="python-docx",
        page_count=None,
        blocks=blocks,
        quality_report=_build_generic_quality_report(
            blocks=blocks,
            source_format="docx",
        ),
    )


def parse_image(path: Path) -> ParsedDocument:
    """Parse an image file using RapidOCR to extract text content.

    Supports PNG, JPG, JPEG, BMP, TIFF, GIF, WEBP formats.
    Requires the optional dependencies in requirements-ocr.txt:
      rapidocr, onnxruntime
    """
    try:
        from rapidocr import RapidOCR
    except ImportError as exc:
        raise DocumentParseError(
            "Image parsing requires optional dependencies 'rapidocr' and 'onnxruntime'. "
            "Install them via: pip install -r requirements-ocr.txt"
        ) from exc

    # P2: RapidOCR() 初始化失败（如 onnxruntime/模型缺失）也转换为 DocumentParseError
    try:
        ocr = RapidOCR()
    except Exception as exc:
        raise DocumentParseError(
            f"RapidOCR initialization failed (check onnxruntime and model files): {exc}"
        ) from exc

    blocks: list[ParsedBlock] = []
    warnings: list[str] = []

    # P1: OCR 运行时失败应 raise，不产生空 chunks 的坏产物
    try:
        result = ocr(str(path), text_score=_RAPIDOCR_TEXT_SCORE)
        lines = _rapidocr_result_to_lines(result)
    except Exception as exc:
        raise DocumentParseError(
            f"OCR engine failed on {path.name}: {exc}"
        ) from exc

    # OCR 成功但无识别文本时，作为 low_quality 处理
    if lines:
        blocks.append(
            ParsedBlock(
                block_type="paragraph",
                text="\n".join(lines),
                page_start=1,
                page_end=1,
                metadata={"ocr": True},
            )
        )
    else:
        warnings.append("no_ocr_text_extracted")

    return ParsedDocument(
        source_format="image",
        parser_name="rapidocr",
        page_count=1,
        blocks=blocks,
        warnings=warnings,
        quality_report=_build_generic_quality_report(
            blocks=blocks,
            source_format="image",
        ),
    )


def _parse_markdown_blocks(text: str) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    paragraph_lines: list[str] = []
    table_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        paragraph = normalize_space("\n".join(paragraph_lines))
        if paragraph:
            blocks.append(ParsedBlock(block_type="paragraph", text=paragraph))
        paragraph_lines = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            blocks.append(ParsedBlock(block_type="table", text="\n".join(table_lines)))
        table_lines = []

    for line in text.splitlines():
        stripped = line.strip()
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        is_table_line = stripped.startswith("|") and stripped.endswith("|")

        if heading:
            flush_paragraph()
            flush_table()
            blocks.append(
                ParsedBlock(
                    block_type="heading",
                    text=heading.group(2).strip(),
                    metadata={"level": len(heading.group(1))},
                )
            )
            continue

        if is_table_line:
            flush_paragraph()
            table_lines.append(stripped)
            continue

        if not stripped:
            flush_paragraph()
            flush_table()
            continue

        flush_table()
        paragraph_lines.append(stripped)

    flush_paragraph()
    flush_table()
    return blocks


def _split_paragraphs(text: str) -> list[str]:
    paragraphs = re.split(r"(?:\r?\n\s*){2,}", text)
    return [normalize_space(paragraph) for paragraph in paragraphs if normalize_space(paragraph)]


def _parsed_document_has_text_blocks(document: ParsedDocument) -> bool:
    return any(
        block.block_type != "heading" and normalize_space(block.text)
        for block in document.blocks
    )


def _build_generic_quality_report(
    *,
    blocks: list[ParsedBlock],
    source_format: str,
    fallback_parser: str | None = None,
) -> dict[str, Any]:
    text_blocks = [
        block for block in blocks if block.block_type != "heading" and normalize_space(block.text)
    ]
    char_count = sum(len(normalize_space(block.text)) for block in text_blocks)
    table_count = sum(1 for block in blocks if block.block_type == "table")
    has_page_numbers = any(block.page_start is not None for block in blocks)
    reason_codes: list[str] = []
    if not text_blocks:
        reason_codes.append("no_text_blocks")
    elif char_count < _MIN_TRUSTED_PDF_TEXT_CHARS:
        reason_codes.append("text_too_short")

    score = 100.0
    if "no_text_blocks" in reason_codes:
        score -= 70.0
    if "text_too_short" in reason_codes:
        score -= 30.0
    if source_format == "pdf" and not has_page_numbers:
        score -= 10.0
    score = max(0.0, min(100.0, score))

    return {
        "score": round(score, 2),
        "status": "ok" if not reason_codes else "low_quality",
        "fallback_used": fallback_parser is not None,
        "fallback_parser": fallback_parser,
        "reason_codes": reason_codes,
        "metrics": {
            "char_count": char_count,
            "text_block_count": len(text_blocks),
            "block_count": len(blocks),
            "table_count": table_count,
            "has_page_numbers": has_page_numbers,
        },
    }


def _is_cjk_char(char: str) -> bool:
    codepoint = ord(char)
    return (
        0x3400 <= codepoint <= 0x4DBF
        or 0x4E00 <= codepoint <= 0x9FFF
        or 0xF900 <= codepoint <= 0xFAFF
        or 0x20000 <= codepoint <= 0x2A6DF
        or 0x2A700 <= codepoint <= 0x2B73F
        or 0x2B740 <= codepoint <= 0x2B81F
        or 0x2B820 <= codepoint <= 0x2CEAF
    )


def _is_mojibake_suspect_char(char: str) -> bool:
    codepoint = ord(char)
    return (
        char in _LATIN1_MOJIBAKE_CHARS
        or 0x3040 <= codepoint <= 0x309F
        or 0x30A0 <= codepoint <= 0x30FF
    )


def _rapidocr_result_to_lines(result: Any) -> list[str]:
    if result is None:
        return []

    if hasattr(result, "txts"):
        txts = list(getattr(result, "txts") or [])
        scores = list(getattr(result, "scores") or [])
        return _filter_ocr_lines(txts, scores)

    if isinstance(result, tuple) and result:
        return _rapidocr_result_to_lines(result[0])

    if isinstance(result, list):
        txts: list[str] = []
        scores: list[float | None] = []
        for item in result:
            if isinstance(item, str):
                txts.append(item)
                scores.append(None)
            elif isinstance(item, (list, tuple)) and len(item) >= 2:
                txts.append(str(item[1]))
                score = item[2] if len(item) >= 3 else None
                scores.append(float(score) if isinstance(score, (int, float)) else None)
        return _filter_ocr_lines(txts, scores)

    return []


def _filter_ocr_lines(txts: list[str], scores: list[float | None]) -> list[str]:
    lines: list[str] = []
    for index, raw_text in enumerate(txts):
        score = scores[index] if index < len(scores) else None
        if score is not None and score < _RAPIDOCR_TEXT_SCORE:
            continue
        text = normalize_space(raw_text)
        if text:
            lines.append(text)
    return lines


class _BlockHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[ParsedBlock] = []
        self._ignore_depth = 0
        self._capture_tag: str | None = None
        self._capture_attrs: dict[str, Any] = {}
        self._buffer: list[str] = []
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style"}:
            self._ignore_depth += 1
            return
        if self._ignore_depth:
            return
        if re.fullmatch(r"h[1-6]", tag):
            self._start_capture(tag, {"level": int(tag[1])})
        elif tag in {"p", "li"}:
            self._start_capture(tag, {})
        elif tag == "tr":
            self._current_row = []
        elif tag in {"td", "th"}:
            self._current_cell = []

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style"}:
            self._ignore_depth = max(0, self._ignore_depth - 1)
            return
        if self._ignore_depth:
            return
        if tag == self._capture_tag:
            self._finish_capture()
        elif tag in {"td", "th"}:
            cell = normalize_space(" ".join(self._current_cell))
            if cell:
                self._current_row.append(cell)
            self._current_cell = []
        elif tag == "tr":
            if self._current_row:
                self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            self._finish_table()

    def handle_data(self, data: str) -> None:
        if self._ignore_depth:
            return
        if self._capture_tag:
            self._buffer.append(data)
        elif self._current_cell is not None:
            self._current_cell.append(data)

    def _start_capture(self, tag: str, attrs: dict[str, Any]) -> None:
        self._finish_capture()
        self._capture_tag = tag
        self._capture_attrs = attrs
        self._buffer = []

    def _finish_capture(self) -> None:
        if not self._capture_tag:
            return
        text = normalize_space(" ".join(self._buffer))
        if text:
            if re.fullmatch(r"h[1-6]", self._capture_tag):
                self.blocks.append(
                    ParsedBlock(
                        block_type="heading",
                        text=text,
                        metadata={"level": self._capture_attrs["level"]},
                    )
                )
            else:
                self.blocks.append(ParsedBlock(block_type="paragraph", text=text))
        self._capture_tag = None
        self._capture_attrs = {}
        self._buffer = []

    def _finish_table(self) -> None:
        rows = [" | ".join(row) for row in self._table_rows if row]
        if rows:
            self.blocks.append(ParsedBlock(block_type="table", text="\n".join(rows)))
        self._table_rows = []
